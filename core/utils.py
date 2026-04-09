import os
import hashlib
from docx import Document


def get_cache_filename(video_path):
    # create unique hash for file path
    file_hash = hashlib.md5(video_path.encode()).hexdigest()
    return f"cache_{file_hash}.txt"


def save_output(summary, filename):
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        with open(filename, "w", encoding="utf-8") as f:
            f.write(summary)

    elif ext == ".docx":
        doc = Document()
        doc.add_heading("Video Summary", 0)

        def clean_text(text):
        # Remove markdown-like syntax
            text = text.replace("**", "")
            text = text.replace("---", "")
            return text.strip()
        
        
        for line in summary.split("\n"):
            cleaned = clean_text(line)
            if cleaned:  # avoid empty lines
                doc.add_paragraph(cleaned)
        doc.save(filename)

    else:
        print("Unsupported file format. Use .txt or .docx")
        return

    print(f"\nSaved summary to {filename}")


def split_text(text, max_words=500):
    words = text.split()
    chunks = []

    for i in range(0, len(words), max_words):
        chunk = " ".join(words[i:i + max_words])
        chunks.append(chunk)

    return chunks